import streamlit as st
import requests
from datetime import datetime
import json
import os
import io
from docx import Document
import base64
import unicodedata

def remove_accents(input_str):
    """Remove acentos de uma string."""
    nfkd_form = unicodedata.normalize('NFKD', input_str)
    return "".join([c for c in nfkd_form if not unicodedata.combining(c)])

def get_base64(file_path):
    """Converte o conteúdo de um arquivo em base64."""
    with open(file_path, "rb") as f:
        data = f.read()
    return base64.b64encode(data).decode()

# Caminhos para os arquivos locais
icon_path = "images/favicon.ico"         # Imagem do ícone (favicon)
background_path = "images/background2.jpg" # Imagem de fundo

# Converte as imagens para base64
icon_base64 = get_base64(icon_path)
bg_base64 = get_base64(background_path)

# Configura a página (essa função deve ser chamada antes de qualquer outro comando do Streamlit)
st.set_page_config(
    page_title="Controle Musical",
    page_icon=f"data:image/x-icon;base64,{icon_base64}"
)

# Injetando CSS para definir a imagem de fundo
st.markdown(
    f"""
    <style>
    .stApp {{
        background: linear-gradient(rgba(0, 0, 0, 0.5), rgba(0, 0, 0, 0.5)),
                    url("data:image/jpeg;base64,{bg_base64}");
        background-size: cover;
        background-position: center;
        background-repeat: no-repeat;
        background-attachment: fixed;
    }}
    </style>
    """,
    unsafe_allow_html=True
)

# Arquivo para persistir as modificações (JSON)
DATA_FILE = "modifications.json"

def load_modifications():
    """Carrega as modificações do arquivo JSON, se existir."""
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE, "r", encoding="utf-8") as f:
            try:
                return json.load(f)
            except json.JSONDecodeError:
                return []
    else:
        return []

def save_modifications(modifications):
    """Salva a lista de modificações no arquivo JSON."""
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(modifications, f, ensure_ascii=False, indent=4)

# Ao iniciar o app, carrega as modificações persistentes
if "modifications" not in st.session_state:
    st.session_state["modifications"] = load_modifications()

def get_local_lyrics(artist: str, title: str):
    """
    Procura a letra da música em um banco de dados local (arquivo JSON).
    O arquivo 'local_lyrics.json' deve conter uma lista de registros no formato:
      [
         {"artist": "Nome do Artista", "title": "Nome da Música", "lyrics": "Letra..."},
         ...
      ]
    A busca é feita de forma case-insensitive e sem considerar acentos.
    """
    local_db_file = "local_lyrics.json"
    if os.path.exists(local_db_file):
        with open(local_db_file, "r", encoding="utf-8") as f:
            try:
                data = json.load(f)
                if not isinstance(data, list):
                    st.warning("O arquivo local_lyrics.json não contém uma lista de registros.")
                    return None
                # Normaliza os valores de busca para remover acentos e colocar em minúsculas
                search_artist = remove_accents(artist.strip().lower())
                search_title = remove_accents(title.strip().lower())
                for record in data:
                    if not isinstance(record, dict):
                        continue  # ignora registros inválidos
                    record_artist = remove_accents(record.get("artist", "").strip().lower())
                    record_title = remove_accents(record.get("title", "").strip().lower())
                    if record_artist == search_artist and record_title == search_title:
                        return record.get("lyrics")
            except json.JSONDecodeError:
                st.error("Erro ao decodificar o arquivo local_lyrics.json.")
                return None
    return None

def get_lyrics(artist: str, title: str):
    """
    Busca a letra da música usando a API Lyrics.ovh.
    Se não encontrar na API, procura no banco de dados local.
    Retorna a letra se encontrada ou None caso contrário.
    """
    url = f"https://api.lyrics.ovh/v1/{artist}/{title}"
    try:
        response = requests.get(url)
        if response.status_code == 200:
            data = response.json()
            if isinstance(data, dict):
                lyrics = data.get("lyrics")
                if lyrics:
                    return lyrics
    except Exception:
        pass
    # Se não encontrou na API, procura no banco de dados local (com busca sem acentos)
    return get_local_lyrics(artist, title)

def singer_page():
    st.title("Buscar Letra")
    st.write("Preencha os campos abaixo para buscar a letra da música.")

    if "last_lyrics" not in st.session_state:
        st.session_state.last_lyrics = None

    # --- Formulário de Busca ---
    with st.form(key="search_form"):
        artist = st.text_input("Nome do Artista", key="artist_input")
        title = st.text_input("Nome da Música", key="title_input")
        search_button = st.form_submit_button("Buscar Letra")
    
    if search_button:
        if artist.strip() and title.strip():
            st.session_state.last_artist = artist.strip()
            st.session_state.last_title = title.strip()
            lyrics = get_lyrics(artist, title)
            st.session_state.last_lyrics = lyrics  # Pode ser None se não encontrado
        else:
            st.error("Por favor, preencha os campos 'Nome do Artista' e 'Nome da Música'.")

    # --- Exibe resultados da busca para edição ou inserção manual ---
    if "last_artist" in st.session_state and "last_title" in st.session_state:
        st.write("**Artista:**", st.session_state.last_artist)
        st.write("**Música:**", st.session_state.last_title)

        default_lyrics = st.session_state.last_lyrics if st.session_state.last_lyrics is not None else ""

        if default_lyrics:
            st.info("Letra encontrada. Você pode editá-la abaixo:")
        else:
            st.info("Letra não encontrada. Insira a letra manualmente abaixo:")

        # --- Formulário de Edição/Inserção da Letra ---
        with st.form(key="edit_form"):
            edited_lyrics = st.text_area("Letra (editável)", value=default_lyrics, key="edited_lyrics", height=300)
            editor_name = st.text_input("Nome", key="editor_name")
            save_button = st.form_submit_button("Salvar Alteração")
        
        if save_button:
            if edited_lyrics.strip() == "":
                st.error("A letra não pode estar vazia!")
            elif editor_name.strip() == "":
                st.error("Por favor, preencha o nome de quem está editando!")
            else:
                modification = {
                    "artist": st.session_state.last_artist,
                    "title": st.session_state.last_title,
                    "lyrics": edited_lyrics,
                    "editor": editor_name.strip(),
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                }
                st.session_state.modifications.append(modification)
                save_modifications(st.session_state.modifications)  # Persiste a alteração
                st.success("Alteração salva com sucesso!")
                # Remove os dados da busca para reiniciar o ciclo
                del st.session_state.last_artist
                del st.session_state.last_title
                del st.session_state.last_lyrics

def admin_page():
    st.title("Letras Salvas")
    if st.session_state.modifications:
        # Percorre cada modificação e exibe seus dados
        for idx, mod in enumerate(st.session_state.modifications):
            st.subheader(f"{idx+1}. {mod['artist']} - {mod['title']}")
            st.write("**Editado por:**", mod["editor"])
            st.text_area("Letra:", mod["lyrics"], height=200, key=f"lyrics_{idx}")
            st.write("**Salvo em:**", mod["timestamp"])
            st.markdown("---")
                    
            # Cria duas colunas para os botões Exportar e Excluir
            col_export, col_delete, *_ = st.columns([3.2, 3])
            with col_export:
                # Cria o documento Word para essa música
                doc = Document()
                doc.add_heading(f"{mod['artist']} - {mod['title']}", level=1)
                doc.add_paragraph(f"Editado por: {mod['editor']}")
                doc.add_paragraph(mod['lyrics'])
                doc.add_paragraph(f"Salvo em: {mod['timestamp']}")
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                st.download_button(
                    label="Exportar para Word",
                    data=doc_io,
                    file_name=f"{mod['artist']}_{mod['title']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_{idx}"
                )
            with col_delete:
                if st.button("Apagar Letra", key=f"delete_{idx}"):
                    st.session_state.modifications.pop(idx)
                    save_modifications(st.session_state.modifications)
                    st.rerun()
            st.markdown("---")
    else:
        st.info("Nenhuma alteração salva até o momento.")

# --- Navegação via Sidebar ---
page = st.sidebar.radio("Navegue entre as páginas", ["Letras", "Consulta"])

if page == "Letras":
    singer_page()
elif page == "Consulta":
    admin_page()